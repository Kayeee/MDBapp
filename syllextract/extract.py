from subprocess import Popen, PIPE

import re
import docx
import sys
import subprocess
import os


def docx_par_to_text(doc):
    pars = doc.paragraphs
    total = ""

    for par in pars:
        total += par.text

    return total

def check_row_for_keywords(row):
    cells = row.cells
    headers = {}

    for index, cell in enumerate(cells):
        if re.search(r'date', cell.text, re.I):
            headers["datesIndex"] = index
        elif re.search(r'assignment', cell.text, re.I):
            headers["assignmentsIndex"] = index
        elif re.search(r'quiz test', cell.text, re.I):
            headers["quiz_testIndex"] = index

    return headers

def get_events_from_tables(doc):
    tables = doc.tables
    events = []

    for table in tables:
        #headers will be a dict of {"date": columnIndex, "assignment": columnIndex, "text/quiz": columnIndex}
        headers = check_row_for_keywords(table.rows[0])

        if len(headers) < 2:
            continue

        date_column = table.columns[headers["datesIndex"]]
        event_column = table.columns[headers["assignmentsIndex"]]

        events = []
        for date_cell, event_cell in zip(date_column.cells, event_column.cells):
            if not (date_cell.text == "" and event_cell.text == ""):
                events.append((date_cell.text, event_cell.text))

    return events

def search_for_class_name(text):
    class_name = re.search(r'([A-Z][A-Z][A-Z] [0-9][0-9][0-9])', text)
    if class_name:
        print class_name.groups(1)[0]
        return class_name.groups(1)[0]

def search_for_instructor(text):
    instructor = re.search(r'instructor.*: (.*)\n', text, re.I)
    if instructor:
        print instructor.groups(1)[0]
        return instructor.groups(1)[0]

def search_for_instructor_email(text):
    instructor_email = re.search(r'e-*mail: (.*)\n', text, re.I)
    if instructor_email:
        print instructor_email.groups(1)[0]
        return instructor_email.groups(1)[0]

def search_params_in_par(text):
    found_params = {}

    class_name = search_for_class_name(text)
    if class_name:
        found_params["class_name"] = class_name

    instructor = search_for_instructor(text)
    if instructor:
        found_params["instructor"] = instructor

    instructor_email = search_for_instructor_email(text)
    if instructor_email:
        found_params["instructor_email"] = instructor_email

    return found_params


def search_params_in_tables(found_params, doc):
    table_found_params = {}

    for table in doc.tables:
        for row in table.rows:
            row_text = ""

            for cell in row.cells:
                row_text += cell.text
            row_text += "\n"

            table_found_params.update(search_params_in_par(row_text).copy())

    #combine param dicts
    for key in table_found_params:
        found_params[key] = table_found_params[key]

    return found_params


def main(passed_file):
    #passsed_file is of type InMemoryUploadedFile (see django docs)
    file_path = passed_file.name

    if file_path[-4:] == ".doc":
        subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file_path])

    elif file_path[-5:] == ".docx":
        d = docx.Document(passed_file)

        #get parameters
        parText = docx_par_to_text(d)
        found_params = search_params_in_par(parText)
        found_params = search_params_in_tables(found_params, d)

        print found_params

        #get events
        events = get_events_from_tables(d)
        for event in events:
            print event

        found_params["events"] = events

        return found_params

    elif file_path[-4] == ".pdf":
        pass


if __name__ == '__main__':
    #open document
    main(sys.argv[1])
